import fitz  # PyMuPDF for PDFs
from docx import Document
import re
from collections import Counter
import pandas as pd
import os

# ----------- Extract text from PDF -----------
def extract_text_from_pdf(file_path):
    text = ""
    doc = fitz.open(file_path)
    for page in doc:
        text += page.get_text()
    return text

# ----------- Extract text from DOCX -----------
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])

# ----------- Common wrapper for Streamlit -----------
def extract_text(uploaded_file, file_type):
    try:
        if file_type == "pdf":
            # Load PDF from in-memory file
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            return ''.join([page.get_text() for page in doc])
        elif file_type == "docx":
            document = Document(uploaded_file)
            return '\n'.join([para.text for para in document.paragraphs])
        else:
            return None
    except Exception as e:
        print(f"❌ Error extracting text: {e}")
        return None


# ----------- Extract keywords from text -----------
def extract_keywords(text, num_keywords=30):
    text = text.lower()
    text = re.sub(r'\S+@\S+', '', text)  # Remove emails
    text = re.sub(r'http\S+|www\S+|linkedin\S+|github\S+', '', text)  # Remove URLs
    text = re.sub(r'\d+', '', text)  # Remove numbers

    words = re.findall(r'\b[a-z]{3,}\b', text)

    stop_words = set([
        "the", "and", "for", "with", "that", "this", "from", "are", "was", 
        "you", "your", "but", "not", "have", "has", "been", "can", "will",
        "had", "they", "their", "what", "when", "how", "who", "why", "where",
        "all", "any", "she", "him", "her", "its", "also", "our", "more", "such",
        "com", "present", "linkedin", "github", "utm", "source", "share", 
        "via", "overview", "tab", "from", "to", "content", "android", "app",
        "profile", "resume", "limited", "student", "aspiring", "working",
        "passionate", "building", "solutions", "abhishek", "india"
    ])

    filtered = [word for word in words if word not in stop_words]
    return [w for w, _ in Counter(filtered).most_common(num_keywords)]

# ----------- (Optional) Save keywords to CSV -----------
def save_keywords_to_csv(keywords, output_file="keywords.csv"):
    df = pd.DataFrame(keywords, columns=['Keyword'])
    df.to_csv(output_file, index=False)

# ----------- Optional: Manual test runner -----------
def process_file(file_path, file_type, output_csv='keywords.csv', num_keywords=30):
    if not os.path.exists(file_path):
        print("❌ File not found.")
        return

    print("🔍 Reading file:", file_path)
    text = extract_text(file_path, file_type)

    if text:
        keywords = extract_keywords(text, num_keywords=num_keywords)
        save_keywords_to_csv(keywords, output_csv)
        print(f"✅ Keywords saved to {output_csv}")
    else:
        print("⚠️ No text extracted.")

# ----------- Run when called directly -----------
if __name__ == "__main__":
    file_path = "My Resume.pdf"
    file_type = "pdf"
    output_csv = "keywords.csv"

    print("🚀 Starting keyword extraction...\n")
    process_file(file_path, file_type, output_csv)
    print("\n✅ Done.")